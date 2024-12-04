import fitz
import logging
from typing import Dict, Optional
import traceback
from utils.logging_config import setup_logging
from docx import Document
import io

# Initialize logging configuration
setup_logging()
logger = logging.getLogger('pdf_processor')

class DocumentProcessingError(Exception):
    """Base exception for document processing errors."""
    pass

class PDFProcessingError(DocumentProcessingError):
    """Custom exception for PDF processing errors."""
    pass

class DOCXProcessingError(DocumentProcessingError):
    """Custom exception for DOCX processing errors."""
    pass

class DocumentProcessor:
    """Base class for document processing."""
    
    @staticmethod
    def extract_sections(text: str) -> Dict[str, str]:
        """
        Extract key resume sections with improved detection and error handling
        
        Args:
            text: Raw text from document
            
        Returns:
            Dict containing extracted sections
        """
        if not isinstance(text, str):
            logger.error(f"Invalid input type for text: {type(text)}")
            raise TypeError("Input text must be a string")
            
        if not text.strip():
            logger.error("Empty text provided for section extraction")
            raise ValueError("No text provided for section extraction")

        try:
            sections = {
                'experience': '',
                'education': '',
                'skills': ''
            }
            
            # Enhanced section detection patterns
            section_keywords = {
                'experience': [
                    'experience', 'work history', 'employment', 'work experience'
                ],
                'education': [
                    'education', 'academic background', 'qualifications',
                    'academic history'
                ],
                'skills': [
                    'skills', 'technical skills', 'core competencies', 'expertise'
                ]
            }
            
            lines = text.split('\n')
            current_section = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                lower_line = line.lower()
                
                # Check for section headers
                for section, keywords in section_keywords.items():
                    if (any(keyword in lower_line for keyword in keywords) and
                            not line.strip().islower()):
                        current_section = section
                        break
                        
                # Add content to current section
                if current_section and line:
                    # Don't add the section header itself
                    if not any(
                        keyword in lower_line 
                        for keyword in section_keywords[current_section]
                    ):
                        sections[current_section] = (
                            sections[current_section] + line + '\n'
                        )
            
            # Clean up trailing whitespace
            for section in sections:
                sections[section] = sections[section].strip()
            
            logger.debug(f"Extracted sections: {list(sections.keys())}")
            return sections
            
        except Exception as e:
            logger.error(f"Failed to extract sections: {str(e)}")
            logger.debug(f"Section extraction error details: {traceback.format_exc()}")
            raise RuntimeError(f"Failed to extract resume sections: {str(e)}") from e

class PDFProcessor(DocumentProcessor):
    """Processor for PDF documents."""
    
    @staticmethod
    def extract_text(pdf_file) -> str:
        """
        Extract text from uploaded PDF with enhanced error handling
        
        Args:
            pdf_file: StreamIO object containing PDF data
            
        Returns:
            str: Extracted text from PDF
            
        Raises:
            PDFProcessingError: If PDF is corrupted or cannot be processed
        """
        if not pdf_file:
            logger.error("No PDF file provided")
            raise PDFProcessingError("No PDF file was provided")

        try:
            # Validate file size
            pdf_file.seek(0, 2)  # Seek to end
            file_size = pdf_file.tell()
            pdf_file.seek(0)  # Reset to beginning
            
            # 100MB limit
            if file_size > 100 * 1024 * 1024:
                logger.error(f"PDF file too large: {file_size / (1024*1024):.2f}MB")
                raise PDFProcessingError("PDF file size exceeds 100MB limit")

            doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
            
            if doc.page_count == 0:
                logger.error("PDF document contains no pages")
                raise PDFProcessingError("The PDF document contains no pages")
                
            text = ""
            for page_num, page in enumerate(doc):
                try:
                    page_text = page.get_text()
                    text += page_text
                    logger.debug(f"Successfully extracted text from page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                    continue  # Continue with next page if one fails
            
            if not text.strip():
                logger.warning("Extracted text is empty")
                raise PDFProcessingError("No readable text found in the PDF")
                
            return text.strip()
            
        except fitz.FileDataError as e:
            logger.error(f"Invalid or corrupted PDF file: {str(e)}")
            raise PDFProcessingError("The PDF file appears to be corrupted or invalid") from e
        except MemoryError as e:
            logger.error(f"Memory error while processing PDF: {str(e)}")
            raise PDFProcessingError("Not enough memory to process the PDF file") from e
        except Exception as e:
            logger.error(f"Failed to process PDF: {str(e)}")
            logger.debug(f"PDF processing error details: {traceback.format_exc()}")
            raise PDFProcessingError("Unable to process PDF file. Please ensure it is not corrupted.") from e
        finally:
            if 'doc' in locals():
                doc.close()

class DOCXProcessor(DocumentProcessor):
    """Processor for DOCX documents."""
    
    @staticmethod
    def extract_text(docx_file) -> str:
        """
        Extract text from uploaded DOCX with enhanced error handling
        
        Args:
            docx_file: StreamIO object containing DOCX data
            
        Returns:
            str: Extracted text from DOCX
            
        Raises:
            DOCXProcessingError: If DOCX is corrupted or cannot be processed
        """
        if not docx_file:
            logger.error("No DOCX file provided")
            raise DOCXProcessingError("No DOCX file was provided")

        try:
            # Validate file size
            docx_file.seek(0, 2)  # Seek to end
            file_size = docx_file.tell()
            docx_file.seek(0)  # Reset to beginning
            
            # 100MB limit
            if file_size > 100 * 1024 * 1024:
                logger.error(f"DOCX file too large: {file_size / (1024*1024):.2f}MB")
                raise DOCXProcessingError("DOCX file size exceeds 100MB limit")

            # Load the document
            doc = Document(docx_file)
            
            # Extract text from paragraphs
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
                    
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            # Join all text with newlines
            full_text = '\n'.join(text).strip()
            
            if not full_text:
                logger.warning("Extracted text is empty")
                raise DOCXProcessingError("No readable text found in the DOCX")
                
            return full_text
            
        except Exception as e:
            logger.error(f"Failed to process DOCX: {str(e)}")
            logger.debug(f"DOCX processing error details: {traceback.format_exc()}")
            raise DOCXProcessingError("Unable to process DOCX file. Please ensure it is not corrupted.") from e
